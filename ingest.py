# ingest.py (Refactored for your stack)

import os
import fitz  # PyMuPDF
import docx
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS

# --- Configuration ---
DOCS_DIR = "documents"
DB_FAISS_PATH = "vectorstore/db_faiss" # Path to save the FAISS index

def load_docs(directory):
    """
    Loads documents from a directory using PyMuPDF for PDFs and python-docx for DOCX.
    Returns a list of LangChain Document objects.
    """
    docs = []
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if filename.endswith(".pdf"):
            # Use PyMuPDF (fitz) to extract text from PDFs
            doc = fitz.open(file_path)
            text = "".join(page.get_text() for page in doc)
            docs.append(Document(page_content=text, metadata={"source": filename}))
            doc.close()
        elif filename.endswith(".docx"):
            # Use python-docx to extract text from Word documents
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            docs.append(Document(page_content=text, metadata={"source": filename}))
    print(f"Loaded {len(docs)} documents from {directory}")
    return docs

def main():
    """
    Main function to ingest documents, process them, and store them in a FAISS vector store.
    """
    print("🚀 Starting the document ingestion process...")

    # 1. LOAD documents using our custom loader
    documents = load_docs(DOCS_DIR)

    if not documents:
        print("No documents found. Please add your PDF or DOCX files to the 'documents' directory.")
        return

    # 2. SPLIT the loaded documents into smaller chunks
    print("Splitting documents into smaller chunks...")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_documents(documents)
    print(f"✅ Split into {len(chunks)} text chunks.")

    # 3. EMBED the chunks and STORE them in a FAISS vector store
    print("Creating text embeddings using sentence-transformers...")
    # This uses the sentence-transformers library you specified
    embeddings = HuggingFaceEmbeddings(
        model_name="all-MiniLM-L6-v2",
        model_kwargs={'device': 'cpu'}
    )

    print("Creating FAISS vector store...")
    # Create the FAISS index from the chunks and embeddings
    db = FAISS.from_documents(chunks, embeddings)

    # Save the FAISS index locally
    db.save_local(DB_FAISS_PATH)

    print("--------------------------------------------------")
    print(f"✅ Success! Your FAISS knowledge base is ready.")
    print(f"Vector store created and saved at: '{DB_FAISS_PATH}'")
    print("--------------------------------------------------")

if __name__ == "__main__":
    main()